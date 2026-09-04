import streamlit as st
from google import genai
import os
import subprocess
import pypdf
import time
import io
import json
from docx import Document
from docx.shared import Pt, Cm
from datetime import datetime

# Must be the very first Streamlit command
st.set_page_config(page_title="ATS Resume Tailor", page_icon="✨", layout="centered")

# --- Custom CSS for Material/Modern Design ---
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;600;800&display=swap');

/* Hide Streamlit components */
header {visibility: hidden;}
#MainMenu {visibility: hidden;}
footer {visibility: hidden;}

html, body, [class*="css"] {
    font-family: 'Outfit', sans-serif !important;
}

[data-testid="stAppViewContainer"] {
    background-color: #09090b;
    background-image: radial-gradient(circle at 50% 0%, #7f1d1d 0%, transparent 60%);
    color: #f1f5f9;
}

/* Titles */
.main-title {
    text-align: center;
    font-weight: 800;
    background: -webkit-linear-gradient(45deg, #ef4444, #3b82f6);
    -webkit-background-clip: text;
    -webkit-text-fill-color: transparent;
    font-size: 3.5rem;
    margin-bottom: 0;
    padding-bottom: 10px;
}
.sub-title {
    text-align: center;
    color: #94a3b8;
    font-size: 1.2rem;
    font-weight: 300;
    margin-bottom: 2.5rem;
}

/* Primary Button Styling */
div.stButton > button {
    background: linear-gradient(135deg, #ef4444 0%, #2563eb 100%);
    color: white;
    border: none;
    border-radius: 12px;
    padding: 0.5rem 1.5rem;
    font-weight: 600;
    box-shadow: 0 4px 15px rgba(239, 68, 68, 0.4);
    transition: all 0.3s ease;
}
div.stButton > button:hover {
    transform: translateY(-2px);
    box-shadow: 0 6px 20px rgba(239, 68, 68, 0.6);
    color: white;
    border: none;
}

/* Text Area */
.stTextArea textarea {
    background-color: rgba(30, 41, 59, 0.7) !important;
    color: #f8fafc !important;
    border: 1px solid #334155 !important;
    border-radius: 12px !important;
    padding: 15px !important;
    font-size: 1rem !important;
}
.stTextArea textarea:focus {
    border-color: #ef4444 !important;
    box-shadow: 0 0 0 2px rgba(239, 68, 68, 0.2) !important;
}

/* Expanders */
.streamlit-expanderHeader {
    background-color: rgba(30, 41, 59, 0.7) !important;
    border-radius: 8px !important;
    border: 1px solid #334155 !important;
}

/* Markdown Horizontal Rule */
hr {
    border-top: 1px solid #334155;
}
</style>
""", unsafe_allow_html=True)

# --- App Header ---
st.markdown('<h1 class="main-title">ATS Resume Tailor 🕸️</h1>', unsafe_allow_html=True)
st.markdown('<p class="sub-title">Instantly tailor your resume and cover letter with the power of Gemini.</p>', unsafe_allow_html=True)

# --- Job Tracker Helpers ---
TRACKER_FILE = "job_applications.json"

def load_applications():
    if not os.path.exists(TRACKER_FILE):
        return []
    try:
        with open(TRACKER_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                return data
    except Exception:
        pass
    return []

def save_applications(apps):
    with open(TRACKER_FILE, "w", encoding="utf-8") as f:
        json.dump(apps, f, indent=2)

def add_application(company, role, date_str):
    apps = load_applications()
    for app in apps:
        if app.get("company") == company and app.get("role") == role and app.get("date_applied") == date_str:
            return False, "This application is already logged."
    
    apps.insert(0, {
        "company": company,
        "role": role,
        "date_applied": date_str
    })
    save_applications(apps)
    return True, "Application logged successfully."

def delete_application(index):
    apps = load_applications()
    if 0 <= index < len(apps):
        apps.pop(index)
        save_applications(apps)

def get_counts(apps):
    today = datetime.today()
    current_month = today.strftime("%Y-%m")
    month_count = sum(1 for app in apps if app.get("date_applied", "").startswith(current_month))
    return month_count, len(apps)

# --- Job Tracker UI ---
st.markdown(
    """
    <style>
    .tracker-box {
        background-color: rgba(30, 41, 59, 0.4);
        border-radius: 12px;
        padding: 15px;
        margin-bottom: 20px;
        border: 1px solid #334155;
    }
    .tracker-title {
        font-size: 1.1rem;
        font-weight: 600;
        margin-bottom: 5px;
        color: #f8fafc;
    }
    .tracker-stat {
        font-size: 1.5rem;
        font-weight: 800;
        color: #ef4444;
    }
    .tracker-stat-total {
        font-size: 1.5rem;
        font-weight: 800;
        color: #3b82f6;
    }
    </style>
    """, unsafe_allow_html=True
)

apps = load_applications()
month_count, total_count = get_counts(apps)

st.markdown(
    f"""
    <div class="tracker-box">
        <div style="display: flex; justify-content: space-around; text-align: center;">
            <div>
                <div class="tracker-title">Applications this month</div>
                <div class="tracker-stat">{month_count}</div>
            </div>
            <div>
                <div class="tracker-title">Total applications logged</div>
                <div class="tracker-stat-total">{total_count}</div>
            </div>
        </div>
    </div>
    """, unsafe_allow_html=True
)

with st.expander("Recent Applications"):
    if not apps:
        st.write("No applications logged yet.")
    else:
        for i, app in enumerate(apps[:10]):
            col1, col2 = st.columns([5, 1])
            with col1:
                st.write(f"**{app.get('date_applied')}** | {app.get('company')} | {app.get('role')}")
            with col2:
                if st.button("Delete", key=f"del_{i}", use_container_width=True):
                    delete_application(i)
                    st.rerun()

with st.expander("Add Application Manually"):
    with st.form("manual_add"):
        m_company = st.text_input("Company")
        m_role = st.text_input("Role")
        m_date = st.date_input("Date", value=datetime.today())
        if st.form_submit_button("Add Application"):
            if m_company and m_role:
                success, msg = add_application(m_company, m_role, m_date.strftime("%Y-%m-%d"))
                if success:
                    st.success(msg)
                    time.sleep(1)
                    st.rerun()
                else:
                    st.error(msg)
            else:
                st.error("Please fill in company and role.")

with st.expander("Application History"):
    history = {}
    for app in apps:
        date_str = app.get("date_applied", "")
        if len(date_str) >= 7:
            month_key = date_str[:7]
            try:
                dt = datetime.strptime(month_key, "%Y-%m")
                display_month = dt.strftime("%B %Y")
                if display_month not in history:
                    history[display_month] = 0
                history[display_month] += 1
            except:
                pass
    if not history:
        st.write("No history available.")
    else:
        for m, count in history.items():
            st.write(f"**{m}** — {count}")

st.markdown("---")


# Cross-platform LaTeX executable detection
def get_latex_cmd():
    import shutil
    local_win_tectonic = os.path.join(".", "tectonic", "tectonic.exe")
    if os.path.exists(local_win_tectonic):
        return [local_win_tectonic]
    if shutil.which("tectonic"):
        return ["tectonic"]
    return ["pdflatex", "-interaction=nonstopmode"]

# Secure API Key Resolution
def get_api_key():
    try:
        if "GEMINI_API_KEY" in st.secrets:
            return st.secrets["GEMINI_API_KEY"]
    except Exception:
        pass

    if os.environ.get("GEMINI_API_KEY"):
        return os.environ.get("GEMINI_API_KEY")
    # Fallback for local testing or custom deployment
    return st.sidebar.text_input("Gemini API Key", type="password", help="Enter your Gemini API Key")

api_key = get_api_key()

# Initialize Session State
if "generated" not in st.session_state:
    st.session_state.generated = False
if "tailored_tex" not in st.session_state:
    st.session_state.tailored_tex = ""
if "cover_letter" not in st.session_state:
    st.session_state.cover_letter = ""
if "pdf_data" not in st.session_state:
    st.session_state.pdf_data = None
if "resume_filename" not in st.session_state:
    st.session_state.resume_filename = "Adnan_Ahmed_Resume.pdf"
if "cover_filename" not in st.session_state:
    st.session_state.cover_filename = "Adnan_Ahmed_CoverLetter.txt"
if "cover_docx_data" not in st.session_state:
    st.session_state.cover_docx_data = None
if "cover_docx_filename" not in st.session_state:
    st.session_state.cover_docx_filename = "Adnan_Ahmed_CoverLetter.docx"
if "current_company" not in st.session_state:
    st.session_state.current_company = ""
if "current_role" not in st.session_state:
    st.session_state.current_role = ""

# --- Main Logic ---
def call_gemini_with_retry(client, prompt, max_retries=3):
    """Wrapper to handle 429 Rate Limits automatically."""
    for i in range(max_retries):
        try:
            return client.models.generate_content(model='gemini-3.5-flash-lite', contents=prompt)
        except Exception as e:
            if "429" in str(e) and i < max_retries - 1:
                st.toast(f"Rate limit hit. Waiting 20 seconds to retry... (Attempt {i+1}/{max_retries})")
                time.sleep(20)
                continue
            raise e

def extract_job_details(client, jd):
    prompt = f"""
    Analyze the following Job Description and extract the Company Name and a short Job Title (max 3-4 words).
    Return EXACTLY the following format with no other text, markdown, or explanation:
    Company Name|Short Job Title
    
    If the company cannot be confidently identified, output:
    Unknown|Short Job Title
    
    Job Description:
    {jd}
    """
    try:
        response = call_gemini_with_retry(client, prompt)
        text = response.text.strip()
        if "|" in text:
            company, role = text.split("|", 1)
            return company.strip(), role.strip()
    except Exception:
        pass
    return "Unknown", "Role"

def sanitize_filename_part(part):
    import re
    part = re.sub(r'[^a-zA-Z0-9_\-]', '', part.replace(' ', '_'))
    return part

def generate_tailored_skills(client, base_skills_tex, jd):
    prompt = f"""
    You are an expert technical recruiter and ATS resume writer.
    Below is the current "Technical Skills" LaTeX section from my resume and a Job Description (JD).
    
    Job Description:
    {jd}
    
    Current Technical Skills (LaTeX):
    {base_skills_tex}
    
    Task:
    Return ONLY a modified version of the Technical Skills LaTeX code tailored for this job.
    
    CRITICAL INSTRUCTIONS:
    1. You MUST maintain exactly THREE skill categories. You can rename the categories to match the JD.
    2. Prioritize skills in this order: explicit JD requirements, JD desirable skills, closely related tech, useful supporting tech.
    3. Use the exact terminology from the JD (e.g. "Microsoft 365", "Active Directory").
    4. DO NOT add generic soft skills (e.g. communication, teamwork, fast learner).
    5. PRESERVE THE EXACT PHYSICAL SIZE. The current section fits exactly into the 1-page layout. The new section MUST occupy the SAME NUMBER OF RENDERED TEXT LINES. Do NOT add extra vertical spacing, extra bullet points, or explanatory text. Remove lower-priority skills to fit if needed.
    6. Ensure the font size and LaTeX structure (e.g., \\footnotesize, \\\\[1pt]) exactly matches the original.
    7. Return ONLY the raw LaTeX code for this section starting with \\section{{Technical Skills}} and ending with \\resumeSubHeadingListEnd. Do NOT wrap in markdown code blocks if possible, or if you do, only output the code.
    8. ESCAPE LATEX SPECIAL CHARACTERS. For example, C# must be written as C\\#, and & must be written as \\&.
    """
    response = call_gemini_with_retry(client, prompt)
    return response.text

def generate_cover_letter(client, base_tex, jd):
    prompt = f"""
    You are an expert career coach writing a cover letter for me.
    Below is my base resume in LaTeX format and a Job Description (JD) I am applying for.
    
    Job Description:
    {jd}
    
    Base Resume (LaTeX):
    {base_tex}
    
    Task:
    Write a professional, natural Australian-style cover letter. The letter should sound like a real person wrote it.
    
    COVER LETTER STYLE:
    - Tone: confident, straightforward, conversational but professional, concise, grounded, specific.
    - AVOID promotional, corporate, or exaggerated language.
    - Do NOT use phrases like: "I was thrilled to see", "I am deeply excited", "rapid ascent", "leading force", "innovative spirit", "perfect fit", "uniquely positioned", "exact skills needed", "I am confident I can...", "drive meaningful impact", "leverage my skills", "passion for excellence", "fast-paced environment" (unless necessary), "I would be honoured".
    - Do not flatter the employer, claim you are an avid follower, or invent personal enthusiasm/familiarity unless explicitly stated in my resume.
    
    STRUCTURE (250-350 words, 3-5 short paragraphs, NO bullet points):
    - Paragraph 1: State the role being applied for. Briefly establish my current professional background. Give one natural sentence explaining why my background is relevant.
    - Paragraph 2: Connect the strongest relevant professional experience to the actual responsibilities in the JD. Use concrete examples. Do not simply repeat resume bullet points.
    - Paragraph 3: Mention relevant technical/education/project experience where useful. Explain how it complements the professional experience. Only include genuinely relevant information.
    - Optional Paragraph 4: Briefly explain why this particular role makes sense for me. Keep company-specific discussion factual and restrained.
    - Final paragraph: Mention Melbourne location and full Australian work rights. Express interest in discussing the opportunity. Close normally.
    
    WRITING STYLE:
    - Write in plain English. Prefer: "I manage...", "I worked on...", "My experience includes...", "This involved...", "I think that experience would transfer well to...", "The role interests me because...".
    - Avoid: "My extensive expertise uniquely equips me...", "I bring a powerful combination...", "My proven track record demonstrates...".
    - Keep sentences reasonably short. No excessive adjectives. No marketing language. Vary sentence structure naturally.
    
    DO NOT REPEAT THE RESUME:
    - Select 1-2 relevant parts of professional experience, and optionally 1 relevant technical/project example. Explain why they matter to THIS role.
    
    FACTUAL ACCURACY:
    - Use ONLY facts supported by the resume. Do not invent skills, employment, achievements, years of experience, or software usage.
    
    PERSONALISATION:
    - Extract company name and job title from the JD. The company name should appear 1-2 times, job title once near the beginning.
    - If the hiring manager's real name is provided in the JD, use: "Dear [Name],". Otherwise use: "Dear [Company Name] Hiring Team,". Avoid "To Whom It May Concern".
    
    AUSTRALIAN CONTEXT:
    - Use Australian English spelling (e.g. organisation, optimise, behaviour).
    
    Return ONLY the plain text of the cover letter. Do not include markdown formatting markers unless genuinely part of the letter.
    """
    response = call_gemini_with_retry(client, prompt)
    return response.text

def create_docx_cover_letter(letter_text, company_name):
    doc = Document()
    
    # Setup margins and styles
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    # Header
    doc.add_paragraph("Adnan Ahmed")
    doc.add_paragraph("Melbourne, VIC")
    doc.add_paragraph("adnan.techme@gmail.com | +61 424 808 695")
    doc.add_paragraph("") # Spacing
    
    date_str = datetime.today().strftime('%d %B %Y')
    doc.add_paragraph(date_str)
    doc.add_paragraph("") # Spacing
    
    doc.add_paragraph("Hiring Team")
    if company_name and company_name.lower() not in ["unknown", "na", ""]:
        doc.add_paragraph(company_name)
    doc.add_paragraph("Melbourne, VIC")
    doc.add_paragraph("") # Spacing
    
    # Content
    for p in letter_text.split('\n'):
        doc.add_paragraph(p)
        
    f = io.BytesIO()
    doc.save(f)
    return f.getvalue()

# User Input
job_description = st.text_area("Job Description", height=200, placeholder="Paste the job description here...")

col_btn1, col_btn2 = st.columns(2)
with col_btn1:
    resume_clicked = st.button("🚀 Generate Tailored Resume", use_container_width=True)
with col_btn2:
    cover_letter_clicked = st.button("✍️ Generate Cover Letter", use_container_width=True)

if resume_clicked or cover_letter_clicked:
    if not job_description.strip():
        st.error("Please paste the job description.")
    else:
        try:
            with open("main.tex", "r", encoding="utf-8") as f:
                base_tex_content = f.read()
        except FileNotFoundError:
            st.error("main.tex file not found in the current directory.")
            st.stop()
            
        try:
            if not api_key:
                st.error("Gemini API Key is missing. Please set GEMINI_API_KEY in secrets or enter it in the sidebar.")
                st.stop()
            client = genai.Client(api_key=api_key)
            
            progress_bar = st.progress(5, text="🧠 Analyzing Job Description...")
            company, role = extract_job_details(client, job_description)
            st.session_state.current_company = company
            st.session_state.current_role = role
            company_clean = sanitize_filename_part(company)
            role_clean = sanitize_filename_part(role)
            
            if company_clean.lower() in ["unknown", "", "na"]:
                st.session_state.resume_filename = f"Adnan_Ahmed_{role_clean}.pdf"
                st.session_state.cover_filename = f"Adnan_Ahmed_{role_clean}_Cover_Letter.txt"
                st.session_state.cover_docx_filename = f"Adnan_Ahmed_{role_clean}_Cover_Letter.docx"
            else:
                st.session_state.resume_filename = f"Adnan_Ahmed_{company_clean}_{role_clean}.pdf"
                st.session_state.cover_filename = f"Adnan_Ahmed_{company_clean}_Cover_Letter.txt"
                st.session_state.cover_docx_filename = f"Adnan_Ahmed_{company_clean}_Cover_Letter.docx"
                
            # Prevent 429 rate limits by spacing out API requests
            time.sleep(2)
            
            if resume_clicked:
                progress_bar.progress(20, text="📝 Drafting tailored Technical Skills...")
                import re
                match = re.search(r'(\\section\s*\{Technical Skills\}.*?\\resumeSubHeadingListEnd)', base_tex_content, re.DOTALL | re.IGNORECASE)
                if not match:
                    st.error("Could not find Technical Skills section in main.tex")
                    st.stop()
                
                base_skills_tex = match.group(1)
                tailored_skills_tex = generate_tailored_skills(client, base_skills_tex, job_description)
                
                if tailored_skills_tex.startswith("```latex"):
                    tailored_skills_tex = tailored_skills_tex[8:]
                elif tailored_skills_tex.startswith("```tex"):
                    tailored_skills_tex = tailored_skills_tex[6:]
                if tailored_skills_tex.endswith("```"):
                    tailored_skills_tex = tailored_skills_tex[:-3]
                    
                tailored_skills_tex = tailored_skills_tex.strip()
                
                tailored_tex = base_tex_content.replace(base_skills_tex, tailored_skills_tex)
                st.session_state.tailored_tex = tailored_tex
                st.session_state.pdf_data = None
                
                progress_bar.progress(70, text="⚙️ Compiling formatting into a beautiful PDF...")
                with open("tailored_resume.tex", "w", encoding="utf-8") as f:
                    f.write(st.session_state.tailored_tex)
                    
                try:
                    if os.path.exists("tailored_resume.pdf"):
                        os.remove("tailored_resume.pdf")
                        
                    cmd = get_latex_cmd() + ["tailored_resume.tex"]
                    subprocess.run(cmd, check=True, capture_output=True)
                    
                    progress_bar.progress(100, text="✅ Finalizing documents...")
                    with open("tailored_resume.pdf", "rb") as f:
                        st.session_state.pdf_data = f.read()
                except Exception as e:
                    print(f"Compilation error: {e}")
                    st.error(f"LaTeX compilation failed: {e}")

            if cover_letter_clicked:
                progress_bar.progress(50, text="✍️ Writing a highly targeted cover letter...")
                cover_letter = generate_cover_letter(client, base_tex_content, job_description)
                if cover_letter.startswith("```text"):
                    cover_letter = cover_letter[7:]
                elif cover_letter.startswith("```"):
                    cover_letter = cover_letter[3:]
                if cover_letter.endswith("```"):
                    cover_letter = cover_letter[:-3]
                
                st.session_state.cover_letter = cover_letter.strip()
                st.session_state.cover_docx_data = create_docx_cover_letter(st.session_state.cover_letter, company)
                progress_bar.progress(100, text="✅ Finalizing cover letter...")

            st.session_state.generated = True
            st.rerun()
            
        except Exception as e:
            if "429" in str(e):
                st.error("Google Gemini API Rate Limit Exceeded (20 requests per minute). Please wait about a minute and try again!")
            elif "413" in str(e):
                st.error("Payload Too Large (413). The job description or base resume is too large for the API. Please try providing a shorter one.")
            else:
                st.error(f"An error occurred: {e}")

# --- Render Results if Generated ---
if st.session_state.generated:
    st.markdown("---")
    st.markdown("<h3 style='text-align: center; color: #ef4444; margin-bottom: 20px;'>🕷️ Your Documents are Ready!</h3>", unsafe_allow_html=True)
    
    columns = []
    if st.session_state.pdf_data:
        columns.append("resume")
    if st.session_state.cover_letter:
        columns.append("cover_txt")
        columns.append("cover_docx")
        
    if not columns:
        st.warning("No documents were successfully generated.")
    else:
        cols = st.columns(len(columns))
        for i, col_type in enumerate(columns):
            with cols[i]:
                if col_type == "resume":
                    st.download_button(
                        label="📥 Download Resume (.pdf)",
                        data=st.session_state.pdf_data,
                        file_name=st.session_state.resume_filename,
                        mime="application/pdf",
                        use_container_width=True
                    )
                elif col_type == "cover_txt":
                    st.download_button(
                        label="📥 Download Cover Letter (.txt)",
                        data=st.session_state.cover_letter,
                        file_name=st.session_state.cover_filename,
                        mime="text/plain",
                        use_container_width=True
                    )
                elif col_type == "cover_docx":
                    st.download_button(
                        label="📥 Download Cover Letter (.docx)",
                        data=st.session_state.cover_docx_data,
                        file_name=st.session_state.cover_docx_filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                    
    if st.session_state.cover_letter:
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("👁️ Preview Cover Letter"):
            st.write(st.session_state.cover_letter)

    st.markdown("---")
    st.markdown("<h3 style='text-align: center; color: #3b82f6; margin-bottom: 20px;'>📝 Log Application</h3>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center; color: #94a3b8; margin-bottom: 20px;'>Once you've applied to the job, log it in your tracker.</p>", unsafe_allow_html=True)
    
    if st.button("Log Application", use_container_width=True):
        today_str = datetime.today().strftime("%Y-%m-%d")
        comp = st.session_state.get("current_company", "Unknown")
        rol = st.session_state.get("current_role", "Unknown")
        success, msg = add_application(comp, rol, today_str)
        if success:
            st.success(f"Logged application for {rol} at {comp} on {today_str}.")
            time.sleep(1)
            st.rerun()
        else:
            st.warning(msg)
